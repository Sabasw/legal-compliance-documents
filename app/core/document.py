# document.py
import os
from pathlib import Path
import pdfplumber
from docx import Document
from typing import Optional
import easyocr
from PIL import Image
import io
import re
import logging
import hashlib
from app.core.ocr import OCRProcessor

logger = logging.getLogger(__name__)

class DocumentProcessor2:
    def __init__(self):
        self.ocr_processor = OCRProcessor()
        logger.info("Document processor initialized with OCR capabilities")

    def extract_text(self, doc_path: str) -> str:
        """Robust text extraction with format detection, OCR fallback, and error handling"""
        try:
            if not os.path.exists(doc_path):
                raise FileNotFoundError(f"Document not found: {doc_path}")

            # Handle image files
            if doc_path.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
                logger.info(f"Processing image file with OCR: {doc_path}")
                return self.ocr_processor.extract_text_from_image(doc_path)

            # Handle PDF files
            if doc_path.lower().endswith('.pdf'):
                if self.ocr_processor.is_scanned_pdf(doc_path):
                    logger.info(f"Processing scanned PDF with OCR: {doc_path}")
                    return self.ocr_processor.process_scanned_pdf(doc_path)
                return self._extract_pdf(doc_path)

            # Handle other document types
            elif doc_path.lower().endswith('.docx'):
                return self._extract_docx(doc_path)
            elif doc_path.lower().endswith(('.txt', '.md', '.rtf')):
                return self._extract_txt(doc_path)
            else:
                logger.warning(f"Unsupported file format: {doc_path}")
                return ""
        except Exception as e:
            logger.error(f"Extraction failed for {doc_path}: {str(e)}")
            return ""

    @staticmethod
    def _extract_pdf(path: str) -> str:
        """PDF extraction with multiple fallback methods"""
        text = ""

        # Try pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}")

        # Fallback to PyPDF2
        try:
            import PyPDF2
            with open(path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text
        except Exception as e:
            logger.error(f"PyPDF2 failed: {str(e)}")
            return ""

    @staticmethod
    def _extract_docx(path: str) -> str:
        """DOCX extraction with error handling"""
        try:
            from docx import Document
            doc = Document(path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return ""

    @staticmethod
    def _extract_txt(path: str) -> str:
        """Text file extraction with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin1', 'iso-8859-1', 'windows-1252']
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"Failed with encoding {encoding}: {str(e)}")
                continue
        logger.error("Failed to decode text file with all attempted encodings")
        return ""

    @staticmethod
    def document_hash(content: str) -> str:
        """Generate consistent hash for document content"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()


class DocumentProcessor:
    @staticmethod
    def extract_text(doc_path: str) -> str:
        """Robust text extraction with format detection and fallbacks"""
        try:
            if doc_path.endswith(".pdf"):
                return DocumentProcessor._extract_pdf(doc_path)
            elif doc_path.endswith(".docx"):
                return DocumentProcessor._extract_docx(doc_path)
            elif doc_path.endswith((".txt", ".md", ".rtf")):
                return DocumentProcessor._extract_txt(doc_path)
            else:
                logger.warning(f"Unsupported file format: {doc_path}")
                return ""
        except Exception as e:
            logger.error(f"Extraction failed for {doc_path}: {str(e)}")
            return ""

    @staticmethod
    def _extract_pdf(path: str) -> str:
        """PDF extraction with multiple fallback methods"""
        text = ""

        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}")

        try:
            import PyPDF2
            with open(path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text
        except Exception as e:
            logger.error(f"PyPDF2 failed: {str(e)}")
            return ""

    @staticmethod
    def _extract_docx(path: str) -> str:
        """DOCX extraction with error handling"""
        try:
            from docx import Document
            doc = Document(path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return ""

    @staticmethod
    def _extract_txt(path: str) -> str:
        """Text file extraction with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin1', 'iso-8859-1', 'windows-1252']
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"Failed with encoding {encoding}: {str(e)}")
                continue
        logger.error("Failed to decode text file with all attempted encodings")
        return ""

    @staticmethod
    def document_hash(content: str) -> str:
        """Generate consistent hash for document content"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()